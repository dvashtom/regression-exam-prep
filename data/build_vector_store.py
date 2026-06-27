"""
Lightweight Vector Store Builder for RAG
Parses PDFs/DOCX, chunks text, generates embeddings via OpenAI API, stores in FAISS.
No heavy dependencies (no langchain, no torch, no sentence-transformers).
"""

import os
import sys
import json
import pickle
import numpy as np
import faiss
from pathlib import Path
from typing import List, Dict

# PDF parsing
from pypdf import PdfReader
# DOCX parsing
from docx import Document as DocxDocument
# OpenAI for embeddings
from openai import OpenAI
from dotenv import load_dotenv

# Load env
load_dotenv(Path(__file__).parent.parent / ".env")

AI_PROXY_URL = os.getenv("AI_PROXY_URL", "http://localhost:8080/v1")
AI_API_KEY = os.getenv("AI_API_KEY", "sk-placeholder")
EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL", "text-embedding-3-small")


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file."""
    try:
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n\n"
        return text.strip()
    except Exception as e:
        print(f"  Warning: Error reading PDF {Path(file_path).name}: {e}")
        return ""


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file."""
    try:
        doc = DocxDocument(file_path)
        text = ""
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text += row_text + "\n"
        return text.strip()
    except Exception as e:
        print(f"  Warning: Error reading DOCX {Path(file_path).name}: {e}")
        return ""


def chunk_text(text: str, chunk_size: int = 800, overlap: int = 150) -> List[str]:
    """Split text into overlapping chunks."""
    if not text:
        return []
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        if chunk.strip():
            chunks.append(chunk.strip())
        start += chunk_size - overlap
    return chunks


def collect_documents(base_dir: str) -> List[Dict]:
    """Collect all documents, extract text, and chunk them."""
    documents = []
    base_path = Path(base_dir)
    
    # 1. Course book
    course_pdf = base_path / "רגרסיה 2026 (1).pdf"
    if course_pdf.exists():
        print(f"📄 Processing: {course_pdf.name}")
        text = extract_text_from_pdf(str(course_pdf))
        if text:
            chunks = chunk_text(text)
            for i, chunk in enumerate(chunks):
                documents.append({
                    "text": chunk,
                    "source": "חוברת הקורס - רגרסיה 2026",
                    "type": "course_book",
                    "chunk_id": i
                })
            print(f"  -> {len(chunks)} chunks")
    
    # 2. Formula sheet
    formula_pdf = base_path / "דף נוסחאות למבחן 2025 (1).pdf"
    if formula_pdf.exists():
        print(f"📄 Processing: {formula_pdf.name}")
        text = extract_text_from_pdf(str(formula_pdf))
        if text:
            chunks = chunk_text(text, chunk_size=500, overlap=100)
            for i, chunk in enumerate(chunks):
                documents.append({
                    "text": chunk,
                    "source": "דף נוסחאות למבחן 2025",
                    "type": "formula_sheet",
                    "chunk_id": i
                })
            print(f"  -> {len(chunks)} chunks")
    
    # 3. Exam files
    exam_dirs = [
        base_path / "regression_exam_prep" / "exams",
        base_path / "מבחנים"
    ]
    
    for exams_dir in exam_dirs:
        if not exams_dir.exists():
            continue
        print(f"\n📂 Processing exams from: {exams_dir.name}")
        
        # PDFs
        for pdf_file in exams_dir.rglob("*.pdf"):
            text = extract_text_from_pdf(str(pdf_file))
            if text and len(text) > 50:
                chunks = chunk_text(text)
                for i, chunk in enumerate(chunks):
                    documents.append({
                        "text": chunk,
                        "source": pdf_file.name,
                        "type": "exam",
                        "chunk_id": i
                    })
                print(f"  📄 {pdf_file.name} -> {len(chunks)} chunks")
        
        # DOCX (skip temp files)
        for docx_file in exams_dir.rglob("*.docx"):
            if docx_file.name.startswith("~$"):
                continue
            text = extract_text_from_docx(str(docx_file))
            if text and len(text) > 50:
                chunks = chunk_text(text)
                for i, chunk in enumerate(chunks):
                    documents.append({
                        "text": chunk,
                        "source": docx_file.name,
                        "type": "exam",
                        "chunk_id": i
                    })
                print(f"  📄 {docx_file.name} -> {len(chunks)} chunks")
    
    return documents


def generate_embeddings_simple(documents: List[Dict]) -> np.ndarray:
    """
    Generate simple TF-IDF-like embeddings locally (no API needed).
    Uses character n-gram hashing for a lightweight embedding.
    """
    print("\n🧠 Generating local embeddings (hash-based)...")
    
    dim = 384  # embedding dimension
    embeddings = np.zeros((len(documents), dim), dtype=np.float32)
    
    for i, doc in enumerate(documents):
        text = doc["text"]
        # Simple hash-based embedding using character trigrams
        vec = np.zeros(dim, dtype=np.float32)
        # Use word-level features
        words = text.split()
        for word in words:
            # Hash each word to a position
            h = hash(word.lower()) % dim
            vec[h] += 1.0
            # Also hash bigrams
            if len(word) > 2:
                for j in range(len(word) - 2):
                    trigram = word[j:j+3]
                    h2 = hash(trigram) % dim
                    vec[h2] += 0.5
        
        # Normalize
        norm = np.linalg.norm(vec)
        if norm > 0:
            vec = vec / norm
        embeddings[i] = vec
    
    return embeddings


def generate_embeddings_api(documents: List[Dict]) -> np.ndarray:
    """Generate embeddings using OpenAI-compatible API."""
    print("\n🧠 Generating embeddings via API...")
    
    client = OpenAI(base_url=AI_PROXY_URL, api_key=AI_API_KEY)
    
    texts = [doc["text"] for doc in documents]
    batch_size = 20
    all_embeddings = []
    
    for i in range(0, len(texts), batch_size):
        batch = texts[i:i+batch_size]
        print(f"  Processing batch {i//batch_size + 1}/{(len(texts)-1)//batch_size + 1}...")
        try:
            response = client.embeddings.create(
                model=EMBEDDING_MODEL,
                input=batch
            )
            batch_embeddings = [item.embedding for item in response.data]
            all_embeddings.extend(batch_embeddings)
        except Exception as e:
            print(f"  API error: {e}")
            print("  Falling back to local embeddings...")
            return generate_embeddings_simple(documents)
    
    return np.array(all_embeddings, dtype=np.float32)


def build_vector_store(documents: List[Dict], output_path: str, use_api: bool = False):
    """Build FAISS index and save metadata."""
    print(f"\n📊 Building vector store with {len(documents)} chunks...")
    
    # Generate embeddings
    if use_api:
        embeddings = generate_embeddings_api(documents)
    else:
        embeddings = generate_embeddings_simple(documents)
    
    dim = embeddings.shape[1]
    print(f"  Embedding dimension: {dim}")
    
    # Build FAISS index
    index = faiss.IndexFlatIP(dim)  # Inner product (cosine similarity for normalized vectors)
    
    # Normalize for cosine similarity
    faiss.normalize_L2(embeddings)
    index.add(embeddings)
    
    # Save
    os.makedirs(output_path, exist_ok=True)
    faiss.write_index(index, os.path.join(output_path, "index.faiss"))
    
    # Save metadata
    metadata = [{"source": d["source"], "type": d["type"], "text": d["text"]} for d in documents]
    with open(os.path.join(output_path, "metadata.pkl"), "wb") as f:
        pickle.dump(metadata, f)
    
    print(f"\n✅ Vector store saved to: {output_path}")
    print(f"   Index size: {index.ntotal} vectors")
    print(f"   Dimension: {dim}")


def main():
    base_dir = str(Path(__file__).parent.parent.parent)
    output_path = str(Path(__file__).parent / "vector_store")
    
    print("=" * 60)
    print("🚀 RAG Vector Store Builder - רגרסיה לינארית")
    print("=" * 60)
    print(f"\n📁 Base directory: {base_dir}")
    print(f"💾 Output path: {output_path}\n")
    
    # Collect documents
    documents = collect_documents(base_dir)
    
    if not documents:
        print("\n❌ No documents found! Check file paths.")
        return
    
    print(f"\n📚 Total chunks collected: {len(documents)}")
    
    # Build vector store (use local embeddings by default, API if available)
    use_api = os.getenv("USE_API_EMBEDDINGS", "false").lower() == "true"
    build_vector_store(documents, output_path, use_api=use_api)
    
    print("\n" + "=" * 60)
    print("🎉 Done! Vector store is ready for the AI Assistant.")
    print("=" * 60)


if __name__ == "__main__":
    main()